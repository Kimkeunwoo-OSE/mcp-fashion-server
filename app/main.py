from fastapi import FastAPI, File, UploadFile, Request, Form, HTTPException
from fastapi.responses import HTMLResponse, FileResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from typing import List, Optional
import shutil, os, pathlib
from uuid import uuid4
from datetime import datetime
import logging
from pydantic import BaseModel
import sys

sys.path.append(r'C:\Users\kku72\Desktop\MCP_Server')
from image_utils import classify_image, validate_image
from clip_classifier import classify_clothing

# 문서 생성을 위한 패키지
from docx import Document
from docx.shared import Inches

# 루트 디렉토리 설정
BASE_DIR = pathlib.Path(__file__).parent.absolute()
TEMPLATE_DIR = os.path.join(BASE_DIR, "templates")

# FastAPI 앱 설정
app = FastAPI(title="의류 이미지 자동 분석 서비스", 
              description="여성 의류 사진을 자동으로 분석하여 착용컷/디테일컷 분류, 스타일 분석, 워드 문서 생성까지 해주는 서비스")

# 로깅 설정
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


app.mount("/static", StaticFiles(directory=os.path.join(BASE_DIR, "static")), name="static")
templates = Jinja2Templates(directory=TEMPLATE_DIR)

# 업로드 및 출력 폴더 설정 및 생성
UPLOAD_FOLDER = os.path.join(BASE_DIR, "static", "uploads")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "static/outputs")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# 허용되는 파일 확장자 정의
ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg", "gif"}

def allowed_file(filename: str) -> bool:
    """파일 확장자 검증 함수"""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

class ImageResult(BaseModel):
    """이미지 분석 결과를 위한 Pydantic 모델"""
    filename: str
    type: str
    category: str
    description: str

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    """기본 페이지 렌더링"""
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/")
async def root_post(request: Request):
    """폼 제출을 /upload로 리다이렉트"""
    return RedirectResponse(url="/upload", status_code=307)

@app.post("/upload", response_class=HTMLResponse)
async def upload_images(request: Request, files: List[UploadFile] = File(...)):
    """이미지 업로드 및 분석 처리"""
    if not files:
        return templates.TemplateResponse("index.html", {
            "request": request, 
            "error": "파일이 업로드되지 않았습니다."
        })
    
    wear_images, detail_images, descriptions = [], [], []
    
    for file in files:
        try:
            # 파일 유효성 검사
            if not file.filename or not allowed_file(file.filename):
                logger.warning(f"허용되지 않는 파일 형식: {file.filename}")
                continue
                
            # 파일을 고유한 ID로 저장 (중복 방지)
            file_id = str(uuid4()) + os.path.splitext(file.filename)[-1]
            file_path = os.path.join(UPLOAD_FOLDER, file_id)
            
            # 파일 저장
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            # 이미지 유효성 검증
            if not validate_image(file_path):
                logger.error(f"잘못된 이미지 파일: {file.filename}")
                os.remove(file_path)  # 잘못된 파일 삭제
                continue
                
            # 이미지 분류 (착용컷/디테일컷)
            image_type = classify_image(file_path)
            
            try:
                category, style_desc, confidence = classify_clothing(file_path)
            except ValueError as e:
                logger.error(f"값 언패킹 오류: {str(e)}")
                category = "오류"
                style_desc = "분류 과정에서 오류가 발생했습니다."
                confidence = 0.0
            
            # 결과 저장
            descriptions.append({
                "filename": file_id,
                "original_name": file.filename,
                "type": image_type,
                "category": category,
                "description": style_desc
            })
            
            # 착용컷/디테일컷 분류
            if image_type == "wear":
                wear_images.append(file_id)
            else:
                detail_images.append(file_id)
                
        except Exception as e:
            logger.error(f"파일 처리 중 오류: {str(e)}")
            continue
    
    # 처리된 이미지가 있는지 확인
    if not descriptions:
        return templates.TemplateResponse("index.html", {
            "request": request, 
            "error": "처리할 수 있는 이미지가 없습니다. 다시 시도해주세요."
        })
    
    # 미리보기 페이지로 결과 전달
    return templates.TemplateResponse("preview.html", {
        "request": request,
        "wear_images": wear_images,
        "detail_images": detail_images,
        "descriptions": descriptions
    })

@app.post("/generate-document")
async def generate_document(
    request: Request,
    filenames: List[str] = Form(...),
    types: List[str] = Form(...),
    categories: List[str] = Form(...),
    descriptions: List[str] = Form(...)
):
    """분석 결과를 기반으로 워드 문서 생성"""
    try:
        # 파일이 있는지 확인
        if not filenames:
            raise HTTPException(status_code=400, detail="문서 생성을 위한 이미지가 없습니다")

        # Word 문서 생성
        doc = Document()
        doc.add_heading("AI 기반 의류 상세페이지", 0)
        
        # 날짜 정보 추가
        doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("=" * 50)
        
        # 착용컷 섹션
        doc.add_heading("착용컷", level=1)
        wear_added = False
        
        for fname, ftype, category, desc in zip(filenames, types, categories, descriptions):
            if ftype.lower() != "wear":
                continue
                
            wear_added = True
            filepath = os.path.join(UPLOAD_FOLDER, fname)
            
            # 파일 존재 확인
            if not os.path.exists(filepath):
                logger.warning(f"파일을 찾을 수 없음: {fname}")
                continue
                
            try:
                # 이미지와 정보 추가
                doc.add_picture(filepath, width=Inches(4.5))
                doc.add_heading(category, level=2)
                doc.add_paragraph(desc)
                doc.add_paragraph("-" * 40)
            except Exception as e:
                logger.error(f"문서에 이미지 추가 중 오류: {str(e)}")
        
        # 착용컷이 없으면 메시지 추가
        if not wear_added:
            doc.add_paragraph("착용컷 이미지가 없습니다.")
            
        # 디테일컷 섹션
        doc.add_heading("디테일컷", level=1)
        detail_added = False
        
        for fname, ftype, category, desc in zip(filenames, types, categories, descriptions):
            if ftype.lower() != "detail":
                continue
                
            detail_added = True
            filepath = os.path.join(UPLOAD_FOLDER, fname)
            
            # 파일 존재 확인
            if not os.path.exists(filepath):
                continue
                
            try:
                # 이미지와 정보 추가
                doc.add_picture(filepath, width=Inches(4.5))
                doc.add_paragraph(f"카테고리: {category}")
                doc.add_paragraph("-" * 40)
            except Exception as e:
                logger.error(f"문서에 이미지 추가 중 오류: {str(e)}")
        
        # 디테일컷이 없으면 메시지 추가
        if not detail_added:
            doc.add_paragraph("디테일컷 이미지가 없습니다.")
        
        # 저장 파일명 생성
        filename = f"의류분석_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join(OUTPUT_FOLDER, filename)
        
        # 문서 저장
        doc.save(output_path)
        
        # 파일 다운로드 응답
        return FileResponse(
            path=output_path, 
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        logger.error(f"문서 생성 중 오류: {str(e)}")
        raise HTTPException(status_code=500, detail=f"문서 생성 중 오류가 발생했습니다: {str(e)}")

@app.get("/preview/{filename}")
async def get_image(filename: str):
    """업로드된 이미지 제공"""
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다")
    return FileResponse(file_path)

@app.on_event("startup")
async def startup_event():
    """서버 시작 시 실행되는 이벤트 핸들러"""
    logger.info("서버가 시작되었습니다. 파일 디렉터리 확인 중...")
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    logger.info(f"업로드 디렉터리: {UPLOAD_FOLDER}")
    logger.info(f"출력 디렉터리: {OUTPUT_FOLDER}")
    
@app.post("/download")
async def download_document(
    request: Request,
    filenames: List[str] = Form(...),
    types: List[str] = Form(...),
    categories: List[str] = Form(...),
    descriptions: List[str] = Form(...)
):
    # generate_document 함수와 동일한 코드 내용
    return await generate_document(request, filenames, types, categories, descriptions)
