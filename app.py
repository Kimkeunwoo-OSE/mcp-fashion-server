import os
from flask import Flask, request, render_template, send_file, flash, redirect, url_for
from werkzeug.utils import secure_filename
from image_utils import classify_image
from clip_classifier import classify_clothing, get_style_description
from docx import Document
from docx.shared import Inches
from datetime import datetime
import pathlib

# 루트 디렉토리 설정 (현재 스크립트 위치 기준)
BASE_DIR = pathlib.Path(__file__).parent.absolute()
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "outputs")

# 디렉토리가 없으면 생성
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["OUTPUT_FOLDER"] = OUTPUT_FOLDER
app.config["SECRET_KEY"] = "your-secret-key-here"  # 플래시 메시지를 위한 시크릿 키 추가

# 허용되는 파일 확장자 정의
ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg", "gif"}

def allowed_file(filename):
    """파일 확장자 검증 함수"""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        # 파일이 요청에 포함되어 있는지 확인
        if "images" not in request.files:
            flash("파일이 없습니다")
            return redirect(request.url)

        files = request.files.getlist("images")
        
        # 파일이 선택되었는지 확인
        if not files or files[0].filename == "":
            flash("선택된 파일이 없습니다")
            return redirect(request.url)

        wear_images, detail_images, descriptions = [], [], []

        for file in files:
            try:
                # 파일 유효성 검사
                if not allowed_file(file.filename):
                    flash(f"{file.filename}은(는) 허용되지 않는 파일 형식입니다")
                    continue

                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
                file.save(filepath)

                # 이미지 분류 시도
                try:
                    image_type = classify_image(filepath)
                    category, _ = classify_clothing(filepath)
                    style_desc = get_style_description(category)
                except Exception as e:
                    # 이미지 처리 중 오류 발생시
                    flash(f"{filename} 처리 중 오류 발생: {str(e)}")
                    # 업로드된 파일 삭제
                    if os.path.exists(filepath):
                        os.remove(filepath)
                    continue

                descriptions.append({
                    "filename": filename,
                    "type": image_type,
                    "category": category,
                    "description": style_desc
                })

                if image_type == "wear":
                    wear_images.append(filename)
                else:
                    detail_images.append(filename)
            
            except Exception as e:
                flash(f"파일 처리 중 오류가 발생했습니다: {str(e)}")
                continue

        # 처리된 이미지가 있는지 확인
        if not descriptions:
            flash("처리할 수 있는 이미지가 없습니다. 다시 시도해주세요.")
            return redirect(request.url)

        return render_template("preview.html",
                               wear_images=wear_images,
                               detail_images=detail_images,
                               descriptions=descriptions)

    return render_template("index.html")


@app.route("/download", methods=["POST"])
def download():
    try:
        data = request.form
        filenames = request.form.getlist("filenames")
        types = request.form.getlist("types")
        descriptions = request.form.getlist("descriptions")

        # 파일이 있는지 확인
        if not filenames:
            flash("다운로드할 이미지가 없습니다")
            return redirect(url_for("index"))

        # Word 문서 생성
        doc = Document()
        doc.add_heading("AI 기반 쇼핑몰 상세페이지", 0)

        for fname, ftype, desc in zip(filenames, types, descriptions):
            filepath = os.path.join(app.config["UPLOAD_FOLDER"], fname)
            
            # 파일 존재 확인
            if not os.path.exists(filepath):
                flash(f"파일 {fname}을(를) 찾을 수 없습니다")
                continue
                
            try:
                doc.add_picture(filepath, width=Inches(4.5))
                doc.add_paragraph(f"[{ftype.upper()}] {fname}")
                doc.add_paragraph(desc)
                doc.add_paragraph("-" * 40)
            except Exception as e:
                flash(f"문서에 {fname} 추가 중 오류 발생: {str(e)}")
                continue

        # 저장 파일명 생성
        filename = f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join(app.config["OUTPUT_FOLDER"], filename)
        
        try:
            doc.save(output_path)
        except Exception as e:
            flash(f"문서 저장 중 오류 발생: {str(e)}")
            return redirect(url_for("index"))

        return send_file(output_path, as_attachment=True)
        
    except Exception as e:
        flash(f"문서 생성 중 오류가 발생했습니다: {str(e)}")
        return redirect(url_for("index"))


@app.route("/uploads/<filename>")
def uploaded_file(filename):
    """업로드된 파일을 제공하는 라우트"""
    return send_file(os.path.join(app.config["UPLOAD_FOLDER"], filename))


if __name__ == "__main__":
    app.run(debug=True)